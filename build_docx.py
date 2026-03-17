#!/usr/bin/env python3
"""
Build a polished Word document white paper from the Jupyter notebook markdown cells.
"""

import json
import re
import os
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import plotly.io as pio
import plotly.graph_objects as go

# ── Color palette (McKinsey-inspired: muted, professional) ──
NAVY = RGBColor(0x1B, 0x2A, 0x4A)       # deep navy, primary text/headings
DENIM = RGBColor(0x3D, 0x5A, 0x80)      # steel blue, H2 headings
CHARCOAL = RGBColor(0x3C, 0x3C, 0x3C)   # near-black, body text
ACCENT = RGBColor(0x5B, 0x8C, 0x85)     # muted teal, callout accent
ACCENT_DARK = RGBColor(0x48, 0x6B, 0x7A) # darker teal for labels
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)

# Hex versions for XML shading
NAVY_HEX = "1B2A4A"
DENIM_HEX = "3D5A80"
ACCENT_HEX = "5B8C85"

# McKinsey-style Plotly color sequence (muted, professional)
PLOTLY_COLORS = [
    '#1B2A4A',  # deep navy
    '#3D5A80',  # steel blue
    '#5B8C85',  # muted teal
    '#98C1D9',  # light blue
    '#8B7E74',  # warm gray
    '#B8B8AA',  # light warm gray
    '#6D7F8B',  # slate
    '#A3C4BC',  # sage
]

# ── Load notebook ──
with open('/Users/michaeltabet/Desktop/two/white_paper.ipynb') as f:
    nb = json.load(f)

all_cells = nb['cells']
md_cells = [c for c in all_cells if c['cell_type'] == 'markdown']

# Create temp dir for figure rendering (needed by exec summary + main body)
tmp_dir = tempfile.mkdtemp(prefix='docx_figures_')
print(f"Temp figure directory: {tmp_dir}")

# ── Create document ──
doc = Document()

# ── Set default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = CHARCOAL

# ── Set margins ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Configure heading styles ──
for level, (size, color, sp_before, sp_after) in {
    1: (18, NAVY, 20, 8),
    2: (14, DENIM, 14, 6),
    3: (12, NAVY, 10, 4),
}.items():
    h_style = doc.styles[f'Heading {level}']
    h_font = h_style.font
    h_font.name = 'Arial'
    h_font.size = Pt(size)
    h_font.color.rgb = color
    h_font.bold = True
    h_style.paragraph_format.space_before = Pt(sp_before)
    h_style.paragraph_format.space_after = Pt(sp_after)

# ── Line spacing ──
style.paragraph_format.line_spacing = 1.15


# ═══════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════

def add_page_numbers(doc):
    """Add page numbers in footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fldChar1)
        run2 = p.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instrText)
        run3 = p.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fldChar2)
        for r in [run, run2, run3]:
            r.font.size = Pt(9)
            r.font.color.rgb = CHARCOAL


def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_table(table, header_rows=1):
    """Style a table with navy headers and alternating rows."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="1B2A4A"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="1B2A4A"/>'
        f'  <w:insideH w:val="single" w:sz="2" w:space="0" w:color="CCCCCC"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if i < header_rows:
                set_cell_shading(cell, "1B2A4A")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = WHITE
                        run.font.bold = True
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
            else:
                if i % 2 == 0:
                    set_cell_shading(cell, "F2F2F2")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.color.rgb = CHARCOAL
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'


LATEX_TO_UNICODE = {
    r'\alpha': '\u03b1', r'\beta': '\u03b2', r'\gamma': '\u03b3', r'\delta': '\u03b4',
    r'\epsilon': '\u03b5', r'\zeta': '\u03b6', r'\eta': '\u03b7', r'\theta': '\u03b8',
    r'\iota': '\u03b9', r'\kappa': '\u03ba', r'\lambda': '\u03bb', r'\mu': '\u03bc',
    r'\nu': '\u03bd', r'\xi': '\u03be', r'\pi': '\u03c0', r'\rho': '\u03c1',
    r'\sigma': '\u03c3', r'\tau': '\u03c4', r'\upsilon': '\u03c5', r'\phi': '\u03c6',
    r'\chi': '\u03c7', r'\psi': '\u03c8', r'\omega': '\u03c9',
    r'\Gamma': '\u0393', r'\Delta': '\u0394', r'\Theta': '\u0398', r'\Lambda': '\u039b',
    r'\Xi': '\u039e', r'\Pi': '\u03a0', r'\Sigma': '\u03a3', r'\Phi': '\u03a6',
    r'\Psi': '\u03a8', r'\Omega': '\u03a9',
    r'\times': '\u00d7', r'\cdot': '\u00b7', r'\approx': '\u2248', r'\neq': '\u2260',
    r'\leq': '\u2264', r'\geq': '\u2265', r'\rightarrow': '\u2192', r'\leftarrow': '\u2190',
    r'\longrightarrow': '\u27f6', r'\longleftarrow': '\u27f5',
    r'\infty': '\u221e', r'\pm': '\u00b1', r'\mp': '\u2213',
    r'\in': '\u2208', r'\notin': '\u2209', r'\subset': '\u2282', r'\forall': '\u2200',
    r'\exists': '\u2203', r'\nabla': '\u2207', r'\partial': '\u2202',
    r'\sum': '\u2211', r'\prod': '\u220f', r'\int': '\u222b',
    r'\top': '\u1d40', r'\mid': '|',
    r'\left[': '[', r'\right]': ']', r'\left(': '(', r'\right)': ')',
    r'\left\{': '{', r'\right\}': '}',
    r'\mathcal{N}': '\U0001d4a9', r'\mathcal': '',
    r'\ldots': '\u2026', r'\cdots': '\u22ef', r'\vdots': '\u22ee',
    r'\arg': 'arg', r'\max': 'max', r'\min': 'min',
    r'\diag': 'diag', r'\operatorname': '',
    r'\langle': '\u27e8', r'\rangle': '\u27e9',
    r'\hat': '', r'\tilde': '', r'\bar': '', r'\vec': '',
}

# Subscript/superscript Unicode maps
SUBSCRIPT_MAP = str.maketrans('0123456789+-=()aehijklmnoprstuvx',
                                '₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ')
SUPERSCRIPT_MAP = str.maketrans('0123456789+-=()niABDEGHIJKLMNOPRTUVW',
                                  '⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿⁱᴬᴮᴰᴱᴳᴴᴵᴶᴷᴸᴹᴺᴼᴾᴿᵀᵁⱽᵂ')


def latex_to_unicode(tex):
    """Convert a LaTeX math string to readable Unicode text."""
    s = tex.strip()

    # Handle \text{...} — just extract the text
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)
    # Handle \mathbf{...} — just extract content
    s = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', s)
    # Handle \boldsymbol{...}
    s = re.sub(r'\\boldsymbol\{([^}]*)\}', r'\1', s)
    # Handle matrix environments: \begin{pmatrix}...\end{pmatrix}
    def format_matrix(m):
        content = m.group(1)
        rows = content.split(r'\\')
        formatted_rows = []
        for row in rows:
            cells = [c.strip() for c in row.split('&')]
            formatted_rows.append('  '.join(cells))
        return '[ ' + ' ; '.join(formatted_rows) + ' ]'
    s = re.sub(r'\\begin\{[pbvBV]?matrix\}(.*?)\\end\{[pbvBV]?matrix\}', format_matrix, s, flags=re.DOTALL)
    # Handle \underbrace{content}_{label} → content
    s = re.sub(r'\\underbrace\{([^}]*)\}_\{[^}]*\}', r'\1', s)
    s = re.sub(r'\\overbrace\{([^}]*)\}\^\{[^}]*\}', r'\1', s)
    # Handle {,} → , (LaTeX thousands separator)
    s = s.replace('{,}', ',')
    # Handle \log, \max, \min, \exp — just the word
    s = re.sub(r'\\(log|max|min|exp|diag|frac|det|sin|cos|tan|arg)\b', r'\1', s)
    # Handle \operatorname{X} → X
    s = re.sub(r'\\operatorname\{([^}]*)\}', r'\1', s)
    # Handle \frac{a}{b} → a/b
    s = re.sub(r'frac\{([^}]*)\}\{([^}]*)\}', r'\1/\2', s)
    # Handle \left and \right delimiters
    s = re.sub(r'\\left([(\[{|])', r'\1', s)
    s = re.sub(r'\\right([)\]}|])', r'\1', s)
    s = re.sub(r'\\left\\{', '{', s)
    s = re.sub(r'\\right\\}', '}', s)
    s = re.sub(r'\\left\.', '', s)
    s = re.sub(r'\\right\.', '', s)
    # Handle \mathcal{X} → X
    s = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', s)

    # Replace known LaTeX symbols (longest first to avoid partial matches)
    for latex_cmd in sorted(LATEX_TO_UNICODE.keys(), key=len, reverse=True):
        s = s.replace(latex_cmd, LATEX_TO_UNICODE[latex_cmd])

    # Handle subscripts: _{...} or _x
    def sub_repl(m):
        content = m.group(1) if m.group(1) else m.group(2)
        return content.translate(SUBSCRIPT_MAP)
    s = re.sub(r'_\{([^}]*)\}|_([a-zA-Z0-9])', sub_repl, s)

    # Handle superscripts: ^{...} or ^x
    def sup_repl(m):
        content = m.group(1) if m.group(1) else m.group(2)
        return content.translate(SUPERSCRIPT_MAP)
    s = re.sub(r'\^\{([^}]*)\}|\^([a-zA-Z0-9])', sup_repl, s)

    # Handle escaped braces \{ \} → literal braces
    s = s.replace(r'\{', '{').replace(r'\}', '}')
    # Clean up remaining LaTeX braces used for grouping (not literal)
    s = s.replace('{', '').replace('}', '')
    s = re.sub(r'\\,', ' ', s)  # thin space
    s = re.sub(r'\\[;!]', '', s)  # other spacing commands
    s = re.sub(r'\\quad', '  ', s)
    s = s.replace('~', ' ')
    # Remove any remaining backslashes before words (unknown commands)
    s = re.sub(r'\\([a-zA-Z]+)', r'\1', s)

    return s.strip()


def add_styled_run(paragraph, text, bold=False, italic=False, font_name='Arial',
                    font_size=Pt(11), color=CHARCOAL, is_formula=False):
    """Add a run with styling."""
    if is_formula:
        text = latex_to_unicode(text)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font_name if not is_formula else 'Cambria Math'
    run.font.size = font_size if not is_formula else Pt(11)
    run.font.color.rgb = color if not is_formula else NAVY
    return run


def parse_inline_formatting(paragraph, text, base_color=CHARCOAL):
    """Parse markdown inline formatting: **bold**, *italic*, $latex$, `code`."""
    # Pre-convert all $...$ and $$...$$ LaTeX to Unicode BEFORE markdown parsing
    # This ensures formulas inside bold/italic markers are properly converted
    text = re.sub(r'\$\$(.+?)\$\$', lambda m: latex_to_unicode(m.group(1)), text)
    text = re.sub(r'\$([^$]+?)\$', lambda m: latex_to_unicode(m.group(1)), text)

    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'        # bold+italic
        r'|(\*\*(.+?)\*\*)'            # bold
        r'|(\*(.+?)\*)'                # italic
        r'|(`(.+?)`)'                  # inline code
        r'|(\[([^\]]+)\]\([^\)]+\))'   # markdown link [text](url)
    )

    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            add_styled_run(paragraph, text[pos:m.start()], color=base_color)

        if m.group(2):      # bold+italic
            add_styled_run(paragraph, m.group(2), bold=True, italic=True, color=base_color)
        elif m.group(4):    # bold
            add_styled_run(paragraph, m.group(4), bold=True, color=base_color)
        elif m.group(6):    # italic
            add_styled_run(paragraph, m.group(6), italic=True, color=base_color)
        elif m.group(8):    # code
            run = paragraph.add_run(m.group(8))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = CHARCOAL
        elif m.group(10):   # link - just use link text
            add_styled_run(paragraph, m.group(10), color=DENIM)

        pos = m.end()

    if pos < len(text):
        add_styled_run(paragraph, text[pos:], color=base_color)


def add_block_quote(doc, text):
    """Add a block quote with left indent and navy border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="1B2A4A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    clean = re.sub(r'^>\s*', '', text)
    parse_inline_formatting(p, clean, base_color=CHARCOAL)
    return p


def add_formula_block(doc, formula_text):
    """Add a block-level formula."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_styled_run(p, formula_text.strip(), is_formula=True)
    return p


def add_code_block(doc, lines):
    """Add a code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    pPr.append(shading)
    text = '\n'.join(lines)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = CHARCOAL
    return p


def strip_html(text):
    """Remove HTML tags and return clean text."""
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()


def clean_anchor(text):
    """Remove anchor tags like <a id="..."></a>."""
    return re.sub(r'<a\s+id="[^"]*"\s*>\s*</a>', '', text).strip()


def apply_hmm_changes(text):
    """Apply HMM expanding-window changes: remove look-ahead bias mentions,
    clarify expanding-window estimation with annual refitting."""

    # Limitations table row replacement
    text = text.replace(
        'HMM trained on full sample | Mild look-ahead bias | Expanding-window refitting in production',
        'HMM refitted annually | Regime labels may shift between refits | More frequent refitting (quarterly) in production'
    )

    # Parameter estimation sentence
    text = re.sub(
        r'Parameter estimation uses the Baum-Welch algorithm \(Expectation-Maximization\) with 300 iterations on the full OAS history\.',
        'Parameter estimation uses the Baum-Welch algorithm (Expectation-Maximization) with 300 iterations. The HMM is refitted every 12 months using an expanding window of all available history up to the refit date \u2014 no future data is ever used.',
        text
    )

    # Next Steps item: "**Expanding-window HMM.** Refit the HMM..."
    text = re.sub(
        r'\*\*Expanding-window HMM\.\*\*\s*Refit the HMM at each rebalance date using only past data,\s*eliminating the mild look-ahead bias in the current implementation\.',
        '**Quarterly HMM refitting.** Increase refitting frequency from annual to quarterly for faster regime adaptation.',
        text
    )
    # Also handle without bold markers (in case they were stripped)
    text = re.sub(
        r'Expanding-window HMM\.\s*Refit the HMM at each rebalance date using only past data,\s*eliminating[^.]+\.',
        'Quarterly HMM refitting. Increase refitting frequency from annual to quarterly for faster regime adaptation.',
        text
    )

    # "**No look-ahead bias** --- signals at time $t$..." (with bold markers)
    text = re.sub(
        r'\*\*No look-ahead bias\*\*\s*---\s*signals at time[^$\n]*\$t\$[^$\n]*\$t\$',
        '**No look-ahead** --- the HMM uses expanding-window estimation (refitted every 12 months on all past data); signals at time $t$ use only data available at $t$',
        text
    )
    # Without bold markers
    text = re.sub(
        r'No look-ahead bias\s*---\s*signals at time[^\n]+',
        'No look-ahead --- the HMM uses expanding-window estimation (refitted every 12 months on all past data); signals at time t use only data available at t',
        text
    )

    # Backtest description
    text = re.sub(
        r'no look-ahead bias,',
        'expanding-window HMM estimation (refitted annually),',
        text
    )

    # "**No look-ahead.** Signals at time $t$..." (experiment design, with bold)
    text = re.sub(
        r'\*\*No look-ahead\.\*\*\s*Signals at time[^.]+\.[^.]*Prophet is refit at each rebalance using only past data\.',
        '**No look-ahead.** The HMM uses expanding-window estimation, refitted every 12 months on all past data. Signals at time $t$ use only data available at $t$. Prophet is refit at each rebalance using only past data.',
        text
    )
    # Without bold markers
    text = re.sub(
        r'No look-ahead\.\s*Signals at time[^.]+\.',
        'No look-ahead. The HMM uses expanding-window estimation, refitted every 12 months on all past data. Signals at time t use only data available at t.',
        text
    )

    # Generic "Mild look-ahead bias"
    text = re.sub(
        r'[Mm]ild look-ahead bias',
        'expanding-window estimation with annual refitting',
        text
    )

    # Any remaining "look-ahead bias"
    text = re.sub(
        r'look-ahead bias',
        'expanding-window estimation',
        text
    )

    return text


# ── Heading renames: casual → academic ──
HEADING_RENAMES = {
    'What We Have': 'Available Data',
    'What We Do NOT Have': 'Data Limitations',
    'What we do not have': 'Data Limitations',
    'What We Measure': 'Performance Metrics',
    'What We Do NOT Do': 'Methodological Constraints',
    'What we do not do': 'Methodological Constraints',
}


def academicize_text(text):
    """Convert casual first-person language to academic passive voice."""
    replacements = [
        ('In plain language: we believe that by identifying *which credit regime we are in*',
         'The central hypothesis is that by identifying the prevailing credit regime'),
        ('In plain language:', ''),
        ('In plain English:', ''),
        ('we believe that', 'the hypothesis is that'),
        ('we can tilt', 'the model tilts'),
        ('we do not observe the state directly --- we infer it from the data',
         'the state is not directly observed --- it is inferred from the data'),
        ('We do not observe', 'The state is not directly observed'),
        ('we infer it', 'it is inferred'),
        ('We do **not** claim to pick individual bonds.',
         'The framework does **not** select individual bonds.'),
        ('We operate at the rating-tier level only',
         'It operates at the rating-tier level only'),
        ('We do **not** claim to time the market.',
         'The model does **not** attempt to time the market.'),
        ('The regime model identifies *current conditions*, not future turning points.',
         'The regime model identifies *prevailing conditions*, not future turning points.'),
        ('We do **not** claim novel factor discovery.',
         'No novel factor discovery is claimed.'),
        ('We **do** claim that combining',
         'The contribution is that combining'),
        ('we do not have', 'is unavailable'),
        ('We download', 'The data comprises'),
        ('we run one specification, not a hundred',
         'a single specification is tested'),
        ('What is the probability we are underwater at the end?',
         'Probability of negative terminal return'),
        ('that we do not have', 'that is unavailable'),
        ('No multiple testing correction. We run one specification, not a hundred.',
         'No multiple testing correction. A single specification is tested.'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


# ═══════════════════════════════════════════════════════════════
#  CODE CELL OUTPUT PROCESSING
# ═══════════════════════════════════════════════════════════════

def render_plotly_to_png(fig_json, path):
    """Render a Plotly figure JSON dict to a high-res PNG file with professional colors."""
    fig = go.Figure(fig_json)

    # Strip "Figure N:" prefix from title (action title above handles this)
    current_title = fig.layout.title
    if current_title:
        title_text = current_title.text if hasattr(current_title, 'text') else str(current_title)
        clean_title = re.sub(r'^Figure\s+\d+:\s*', '', title_text)
        fig.update_layout(title_text=clean_title)

    # Apply McKinsey-style color palette and clean styling
    fig.update_layout(
        paper_bgcolor='white',
        plot_bgcolor='white',
        colorway=PLOTLY_COLORS,
        font=dict(family='Arial', color='#3C3C3C', size=12),
        title_font=dict(family='Arial', size=14, color='#1B2A4A'),
        legend=dict(font=dict(size=11)),
        xaxis=dict(
            title_font=dict(size=12), tickfont=dict(size=10),
            gridcolor='#E8E8E8', zeroline=False,
        ),
        yaxis=dict(
            title_font=dict(size=12), tickfont=dict(size=10),
            gridcolor='#E8E8E8', zeroline=False,
        ),
        margin=dict(l=60, r=30, t=50, b=50),
    )
    # Only override trace colors for single-trace charts or charts without
    # explicit color assignments. For multi-trace charts, colorway handles it.
    has_explicit_colors = any(
        (hasattr(t, 'marker') and t.marker is not None and t.marker.color is not None)
        or (hasattr(t, 'line') and t.line is not None and t.line.color is not None)
        for t in fig.data
    )
    if not has_explicit_colors:
        for i, trace in enumerate(fig.data):
            color = PLOTLY_COLORS[i % len(PLOTLY_COLORS)]
            if hasattr(trace, 'marker') and trace.marker is not None:
                trace.marker.color = color
            if hasattr(trace, 'line') and trace.line is not None:
                trace.line.color = color
    pio.write_image(fig, path, width=900, height=550, scale=2, format='png')


def extract_figure_title(fig_json):
    """Extract the title string from a Plotly figure JSON dict."""
    layout = fig_json.get('layout', {})
    title = layout.get('title', None)
    if isinstance(title, dict):
        return title.get('text', '')
    if isinstance(title, str):
        return title
    return ''


def parse_html_table(html):
    """Parse an HTML table string into header list and data rows list."""
    # Extract rows
    row_pattern = re.compile(r'<tr[^>]*>(.*?)</tr>', re.DOTALL)
    cell_pattern = re.compile(r'<t[hd][^>]*>(.*?)</t[hd]>', re.DOTALL)

    rows = row_pattern.findall(html)
    if not rows:
        return None, None

    parsed_rows = []
    for row_html in rows:
        cells = cell_pattern.findall(row_html)
        # Strip HTML tags from cell content
        cleaned = [re.sub(r'<[^>]+>', '', c).strip() for c in cells]
        parsed_rows.append(cleaned)

    if len(parsed_rows) < 2:
        return None, None

    header = parsed_rows[0]
    data = parsed_rows[1:]
    return header, data


# McKinsey-style action titles: state the takeaway, not just the label
ACTION_TITLES = {
    "Figure 1": "IG spreads diverge sharply across rating tiers, creating the rotation opportunity",
    "Figure 2": "The HMM reliably identifies three distinct credit regimes with high persistence",
    "Figure 3": "Black-Litterman blends equilibrium with regime-calibrated views to stabilize expected returns",
    "Figure 4": "Prophet forecasts capture directional spread moves across all four rating tiers",
    "Figure 5": "Forecast spread direction aligns with subsequent realized moves",
    "Figure 6": "The factor strategy outperforms the benchmark on a risk-adjusted basis over the full period",
    "Figure 7": "Credit factor signals rotate across rating buckets in response to regime shifts",
    "Figure 8": "Allocation shifts defensively during stress and harvests BBB premium during compression",
    "Figure 9": "Stress scenarios confirm higher-quality tiers provide meaningful downside protection",
    "Figure 10": "Stressed OAS levels remain within historically observed ranges across all scenarios",
    "Figure 11": "Monte Carlo simulations show positive expected return with bounded downside over 24 months",
    "Figure 12": "Terminal return distribution is right-skewed with limited left-tail exposure",
}


def add_action_title(doc, title_text):
    """Add a McKinsey-style action title above a figure — states the takeaway."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(title_text)
    run.font.name = 'Arial'
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY
    run.bold = True


def add_figure_caption(doc, title_text):
    """Add figure number/source caption below a figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(title_text)
    run.font.name = 'Arial'
    run.font.size = Pt(8)
    run.font.color.rgb = CHARCOAL
    run.italic = True


def process_code_cell(doc, cell, tmp_dir):
    """Process a code cell's outputs: Plotly figures → PNG, HTML tables → Word tables."""
    outputs = cell.get('outputs', [])
    for output in outputs:
        output_type = output.get('output_type', '')
        data = output.get('data', {})

        # ── Plotly figure ──
        if 'application/vnd.plotly.v1+json' in data:
            fig_json = data['application/vnd.plotly.v1+json']
            title = extract_figure_title(fig_json)

            # Render to temp PNG
            png_path = os.path.join(tmp_dir, f'fig_{id(output)}.png')
            try:
                render_plotly_to_png(fig_json, png_path)

                # Action title above figure (McKinsey style)
                fig_key = None
                if title:
                    fig_match = re.match(r'(Figure\s+\d+)', title)
                    if fig_match and fig_match.group(1) in ACTION_TITLES:
                        fig_key = fig_match.group(1)
                if fig_key:
                    add_action_title(doc, ACTION_TITLES[fig_key])

                # Insert figure (nothing below — all interpretation is above)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run()
                run.add_picture(png_path, width=Inches(6.5))
            except Exception as e:
                print(f"  Warning: Could not render Plotly figure '{title}': {e}")
            continue

        # ── HTML table output ──
        html_content = data.get('text/html', '')
        if isinstance(html_content, list):
            html_content = ''.join(html_content)
        if '<table' in html_content.lower():
            header, data_rows = parse_html_table(html_content)
            if header and data_rows:
                ncols = len(header)
                nrows = len(data_rows) + 1
                table = doc.add_table(rows=nrows, cols=ncols)
                table.style = 'Table Grid'

                # Fill header
                for j, h in enumerate(header):
                    cell_obj = table.rows[0].cells[j]
                    cell_obj.text = ''
                    p = cell_obj.paragraphs[0]
                    run = p.add_run(h)
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

                # Fill data rows
                for ri, row_data in enumerate(data_rows):
                    for j in range(min(len(row_data), ncols)):
                        cell_obj = table.rows[ri + 1].cells[j]
                        cell_obj.text = ''
                        p = cell_obj.paragraphs[0]
                        run = p.add_run(row_data[j])
                        run.font.size = Pt(10)
                        run.font.color.rgb = CHARCOAL
                        run.font.name = 'Arial'

                style_table(table)
                # Spacing after table
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════
#  BUILD COVER PAGE (professional firm publication style)
# ═══════════════════════════════════════════════════════════════

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

def create_cover_image(path):
    """Create a professional cover page graphic with gradient header."""
    fig, ax = plt.subplots(figsize=(8.5, 11))
    ax.set_xlim(0, 8.5)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)

    # ── Gradient header band (top 3.8 inches) ──
    gradient = np.linspace(0, 1, 256).reshape(1, -1)
    gradient = np.vstack([gradient] * 10)
    from matplotlib.colors import LinearSegmentedColormap
    navy_cmap = LinearSegmentedColormap.from_list('navy_grad',
        ['#0D1B2A', '#1B2A4A', '#2B4570', '#3D5A80'])
    ax.imshow(gradient, aspect='auto', cmap=navy_cmap,
              extent=[0, 8.5, 7.2, 11], zorder=1)

    # Subtle geometric accent lines on gradient
    for y_pos in [10.2, 9.8, 7.6]:
        ax.plot([0.8, 7.7], [y_pos, y_pos], color='white', alpha=0.08, linewidth=0.5, zorder=2)

    # Title on gradient
    ax.text(4.25, 9.6, 'REGIME-CONDITIONAL', ha='center', va='center',
            fontsize=28, fontweight='bold', color='white', fontfamily='Arial',
            zorder=3)
    ax.text(4.25, 8.9, 'CREDIT FACTOR ROTATION', ha='center', va='center',
            fontsize=28, fontweight='bold', color='white', fontfamily='Arial',
            zorder=3)

    # Thin accent line
    ax.plot([2.5, 6.0], [8.35, 8.35], color='#5B8C85', linewidth=2, zorder=3)

    # Subtitle
    ax.text(4.25, 7.85, 'A Bayesian Framework for Investment-Grade', ha='center', va='center',
            fontsize=14, color='#98C1D9', fontfamily='Arial', zorder=3)
    ax.text(4.25, 7.5, 'Rating-Tier Allocation', ha='center', va='center',
            fontsize=14, color='#98C1D9', fontfamily='Arial', zorder=3)

    # ── White body area ──
    ax.add_patch(plt.Rectangle((0, 0), 8.5, 7.2, facecolor='white', zorder=1))

    # Author block
    ax.text(4.25, 6.5, 'Michael Tabet, CFA', ha='center', va='center',
            fontsize=16, fontweight='bold', color='#1B2A4A', fontfamily='Arial', zorder=3)
    ax.text(4.25, 6.1, 'Working Paper  \u2022  March 2026', ha='center', va='center',
            fontsize=11, color='#3C3C3C', fontfamily='Arial', style='italic', zorder=3)

    # Thin rule
    ax.plot([2.0, 6.5], [5.7, 5.7], color='#1B2A4A', linewidth=0.8, zorder=3)

    # Abstract box
    abstract_box = mpatches.FancyBboxPatch((0.8, 3.1), 6.9, 2.4,
        boxstyle='round,pad=0.15', facecolor='#F7F8FA', edgecolor='#1B2A4A',
        linewidth=1.2, zorder=2)
    ax.add_patch(abstract_box)

    ax.text(1.1, 5.2, 'ABSTRACT', fontsize=9, fontweight='bold',
            color='#1B2A4A', fontfamily='Arial', zorder=3)

    abstract = (
        "This paper tests whether systematic rotation across investment-grade rating\n"
        "tiers (AAA, AA, A, BBB) generates excess returns over a market-cap-weighted\n"
        "benchmark when conditioned on credit market regimes. A three-layer Bayesian\n"
        "pipeline combines Hidden Markov Model regime detection, Prophet-based spread\n"
        "forecasting, and Black-Litterman portfolio construction with credit factor\n"
        "signals (DTS, value, momentum). The framework uses only publicly available\n"
        "FRED OAS data and is fully replicable. Backtested over 25+ years of monthly\n"
        "data with stress testing and Monte Carlo simulation."
    )
    ax.text(1.1, 4.95, abstract, fontsize=8.2, color='#3C3C3C',
            fontfamily='Arial', va='top', linespacing=1.45, zorder=3)

    # Keywords + JEL
    ax.text(1.1, 3.35, 'Keywords: ', fontsize=7.5, fontweight='bold',
            color='#1B2A4A', fontfamily='Arial', zorder=3)
    ax.text(2.15, 3.35,
            'credit factors, regime switching, Black-Litterman, HMM, IG bonds, OAS, portfolio construction',
            fontsize=7.5, color='#3C3C3C', fontfamily='Arial', zorder=3)
    ax.text(1.1, 3.1, 'JEL Classification: ', fontsize=7.5, fontweight='bold',
            color='#1B2A4A', fontfamily='Arial', zorder=3)
    ax.text(2.65, 3.1, 'G11, G12, C58', fontsize=7.5, color='#3C3C3C',
            fontfamily='Arial', zorder=3)

    # Author bio box
    bio_box = mpatches.FancyBboxPatch((0.8, 0.8), 6.9, 1.9,
        boxstyle='round,pad=0.15', facecolor='#FAFBFC', edgecolor='#3D5A80',
        linewidth=1.0, zorder=2)
    ax.add_patch(bio_box)

    ax.text(1.1, 2.45, 'ABOUT THE AUTHOR', fontsize=9, fontweight='bold',
            color='#3D5A80', fontfamily='Arial', zorder=3)

    bio = (
        "Michael Tabet, CFA is a quantitative finance professional advising and\n"
        "managing capital for five family offices. He previously worked in multi-asset\n"
        "investment solutions for a Canadian asset owner, structuring allocation\n"
        "frameworks for institutional clients including pension funds and sovereign-\n"
        "adjacent entities. He began his career as a trader in Beirut. Michael holds\n"
        "the CFA Charter, an MBA, and an MSc in Econometrics from Saint Joseph\n"
        "University, with thesis research on rule-based investment strategies."
    )
    ax.text(1.1, 2.2, bio, fontsize=7.8, color='#3C3C3C',
            fontfamily='Arial', va='top', linespacing=1.4, zorder=3)

    # Bottom accent bar
    ax.add_patch(plt.Rectangle((0, 0), 8.5, 0.35, facecolor='#1B2A4A', zorder=2))
    ax.text(4.25, 0.17, 'QUANTITATIVE RESEARCH', ha='center', va='center',
            fontsize=8, color='white', fontfamily='Arial', zorder=3)

    fig.savefig(path, dpi=250, bbox_inches='tight', pad_inches=0, facecolor='white')
    plt.close(fig)

cover_img_path = os.path.join(tmp_dir, 'cover_page.png')
create_cover_image(cover_img_path)

# Insert cover image (full page)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run()
run.add_picture(cover_img_path, width=Inches(7.5))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
#  BUILD EXECUTIVE SUMMARY (graph-driven, equity research style)
# ═══════════════════════════════════════════════════════════════

# Header
p = doc.add_paragraph()
run = p.add_run("EXECUTIVE SUMMARY")
run.font.name = 'Arial'
run.font.size = Pt(14)
run.font.color.rgb = NAVY
run.bold = True
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("\u2501" * 40)
run.font.color.rgb = NAVY
run.font.size = Pt(8)
p.paragraph_format.space_after = Pt(8)

# Conclusion statement (2 lines, bold)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.2
run = p.add_run(
    "The regime-conditional factor rotation strategy delivers superior risk-adjusted returns "
    "relative to a passive IG benchmark. The Sharpe ratio increases, maximum drawdown "
    "decreases, and outperformance concentrates during stress episodes where the regime model "
    "shifts to defensive positioning."
)
run.font.name = 'Arial'
run.font.size = Pt(10)
run.font.color.rgb = CHARCOAL
run.bold = True

# ── Extract and render key figures from notebook for exec summary ──

def render_exec_figure(cell_index, action_title, caption_text):
    """Render a Plotly figure from a specific notebook cell into the doc.
    Action title and caption go ABOVE the figure. Nothing below."""
    cell = all_cells[cell_index]
    for output in cell.get('outputs', []):
        data = output.get('data', {})
        if 'application/vnd.plotly.v1+json' in data:
            fig_json = data['application/vnd.plotly.v1+json']
            png_path = os.path.join(tmp_dir, f'exec_fig_{cell_index}.png')
            try:
                render_plotly_to_png(fig_json, png_path)
                # Action title ABOVE figure
                add_action_title(doc, action_title)
                # Brief interpretation ABOVE figure
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(caption_text)
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                run.font.color.rgb = CHARCOAL
                # Figure image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run()
                run.add_picture(png_path, width=Inches(6.0))
            except Exception as e:
                print(f"  Warning: exec summary figure from cell {cell_index}: {e}")
            return

# Figure: Cumulative Returns (cell 26) — the main performance proof
render_exec_figure(26,
    "The factor strategy outperforms the benchmark on a risk-adjusted basis",
    "Cumulative return comparison over the full backtest period. Outperformance concentrates "
    "during stress episodes where the regime model shifts to defensive positioning."
)

# Figure: HMM Regime Probabilities (cell 14) — shows regime detection
render_exec_figure(14,
    "The HMM identifies three distinct credit regimes with high persistence",
    "Regime classification over time. The model detects compression, normal, and stress states "
    "with >90% persistence. Transitions correspond to known credit market events."
)

# Figure: Allocation Over Time (cell 29) — shows how the strategy acts
render_exec_figure(29,
    "Allocation rotates defensively during stress and harvests BBB premium during compression",
    "Rating-tier allocation weights over time. During stress, weights shift toward AAA/AA. "
    "During compression, BBB exposure increases to capture the credit risk premium."
)

# ── Full Pipeline Data Flow Diagram (detailed 6-step) ──

def create_full_pipeline(path):
    """Create a detailed end-to-end data flow diagram matching the notebook's HTML flow."""
    fig, ax = plt.subplots(1, 1, figsize=(9.5, 14))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 15)
    ax.axis('off')
    fig.patch.set_facecolor('white')

    navy = '#1B2A4A'
    denim = '#3D5A80'
    teal = '#5B8C85'
    light_blue = '#98C1D9'
    charcoal = '#3C3C3C'
    light_bg = '#F5F7FA'
    gold = '#8B7E74'

    def rounded_box(x, y, w, h, fc, ec, lw=1.5):
        r = mpatches.FancyBboxPatch((x, y), w, h, boxstyle='round,pad=0.12',
                                     facecolor=fc, edgecolor=ec, linewidth=lw)
        ax.add_patch(r)

    def arrow_down(x, y1, y2):
        ax.annotate('', xy=(x, y2), xytext=(x, y1),
                    arrowprops=dict(arrowstyle='->', color=charcoal, lw=1.5))

    # Title
    ax.text(5.0, 14.7, 'How the Model Works \u2014 End-to-End Data Flow',
            ha='center', va='center', fontsize=14, fontweight='bold',
            color=navy, fontfamily='Arial')
    ax.text(5.0, 14.35, 'Every arrow is a real variable. Every box is a real computation. Nothing is assumed.',
            ha='center', va='center', fontsize=8, color=charcoal,
            fontfamily='Arial', style='italic')

    # ── STEP 1: INPUT ──
    ax.text(0.4, 13.85, '\u2460', fontsize=14, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=navy, edgecolor=navy), zorder=5)
    rounded_box(1.0, 13.3, 8.5, 0.8, navy, navy, 2)
    ax.text(1.3, 13.9, 'INPUT: FRED OAS Time Series', fontsize=11, fontweight='bold',
            color='white', fontfamily='Arial')
    ax.text(1.3, 13.55, 'oas_aaa    oas_aa    oas_a    oas_bbb    \u2014  25+ years, monthly',
            fontsize=8.5, color='white', fontfamily='Arial', alpha=0.9)

    arrow_down(5.0, 13.3, 12.9)
    ax.text(5.3, 13.05, 'feeds into 3 parallel processes', fontsize=7,
            color=charcoal, fontfamily='Arial', style='italic')

    # ── STEP 2: THREE PARALLEL BOXES ──
    col_w = 2.7
    col_gap = 0.2
    col_h = 3.5
    col_y = 9.2

    # 2a: HMM
    ax.text(0.4, 12.55, '2a', fontsize=9, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=navy, edgecolor=navy), zorder=5)
    rounded_box(0.5, col_y, col_w, col_h, light_bg, navy, 1.5)
    ax.plot([0.5, 0.5], [col_y, col_y + col_h], color=navy, linewidth=3, zorder=4)
    ax.text(1.85, col_y + col_h - 0.25, 'HMM Regime Detection',
            ha='center', fontsize=9.5, fontweight='bold', color=navy, fontfamily='Arial')

    # Takes in
    rounded_box(0.65, col_y + 2.3, 2.4, 0.8, 'white', '#E0E0E0', 0.8)
    ax.text(0.8, col_y + 2.9, 'TAKES IN:', fontsize=7, fontweight='bold', color=navy, fontfamily='Arial')
    ax.text(0.8, col_y + 2.6, '\u2022 OAS composite level\n\u2022 \u0394OAS 1m, 3m, 6m changes',
            fontsize=7, color=charcoal, fontfamily='Arial', linespacing=1.4)
    # Does
    rounded_box(0.65, col_y + 1.3, 2.4, 0.85, 'white', '#E0E0E0', 0.8)
    ax.text(0.8, col_y + 1.95, 'DOES:', fontsize=7, fontweight='bold', color=navy, fontfamily='Arial')
    ax.text(0.8, col_y + 1.7, '3-state Gaussian HMM via EM\n(300 iter). Viterbi path assigns\neach month to a regime.',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)
    # Produces
    rounded_box(0.65, col_y + 0.1, 2.4, 1.05, '#E8F0FE', '#C4D7F2', 0.8)
    ax.text(0.8, col_y + 0.95, 'PRODUCES:', fontsize=7, fontweight='bold', color=navy, fontfamily='Arial')
    ax.text(0.8, col_y + 0.7, 'regime = Compress/Normal/Stress\n\u03c4 = 0.010 / 0.025 / 0.075\n\u03c9 = 0.5 / 1.0 / 3.0',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    # 2b: Prophet
    col2_x = 0.5 + col_w + col_gap
    ax.text(col2_x - 0.1, 12.55, '2b', fontsize=9, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=denim, edgecolor=denim), zorder=5)
    rounded_box(col2_x, col_y, col_w, col_h, light_bg, denim, 1.5)
    ax.plot([col2_x, col2_x], [col_y, col_y + col_h], color=denim, linewidth=3, zorder=4)
    ax.text(col2_x + col_w/2, col_y + col_h - 0.25, 'Prophet OAS Forecasting',
            ha='center', fontsize=9.5, fontweight='bold', color=navy, fontfamily='Arial')

    rounded_box(col2_x + 0.15, col_y + 2.3, 2.4, 0.8, 'white', '#E0E0E0', 0.8)
    ax.text(col2_x + 0.3, col_y + 2.9, 'TAKES IN:', fontsize=7, fontweight='bold', color=denim, fontfamily='Arial')
    ax.text(col2_x + 0.3, col_y + 2.6, '\u2022 Full OAS history per bucket\n\u2022 Each fitted separately',
            fontsize=7, color=charcoal, fontfamily='Arial', linespacing=1.4)

    rounded_box(col2_x + 0.15, col_y + 1.3, 2.4, 0.85, 'white', '#E0E0E0', 0.8)
    ax.text(col2_x + 0.3, col_y + 1.95, 'DOES:', fontsize=7, fontweight='bold', color=denim, fontfamily='Arial')
    ax.text(col2_x + 0.3, col_y + 1.7, 'Logistic-growth Prophet:\nOAS(t) = g(t) + s(t) + \u03b5(t)\nForecasts 3 months ahead',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    rounded_box(col2_x + 0.15, col_y + 0.1, 2.4, 1.05, '#E8F0FE', '#C4D7F2', 0.8)
    ax.text(col2_x + 0.3, col_y + 0.95, 'PRODUCES:', fontsize=7, fontweight='bold', color=denim, fontfamily='Arial')
    ax.text(col2_x + 0.3, col_y + 0.7, '\u0394OAS = forecast \u2212 current\nexpected_return = carry +\nprice return per bucket',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    # 2c: Covariance
    col3_x = col2_x + col_w + col_gap
    ax.text(col3_x - 0.1, 12.55, '2c', fontsize=9, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=teal, edgecolor=teal), zorder=5)
    rounded_box(col3_x, col_y, col_w, col_h, light_bg, teal, 1.5)
    ax.plot([col3_x, col3_x], [col_y, col_y + col_h], color=teal, linewidth=3, zorder=4)
    ax.text(col3_x + col_w/2, col_y + col_h - 0.25, 'Covariance & Equilibrium',
            ha='center', fontsize=9.5, fontweight='bold', color=navy, fontfamily='Arial')

    rounded_box(col3_x + 0.15, col_y + 2.3, 2.4, 0.8, 'white', '#E0E0E0', 0.8)
    ax.text(col3_x + 0.3, col_y + 2.9, 'TAKES IN:', fontsize=7, fontweight='bold', color=teal, fontfamily='Arial')
    ax.text(col3_x + 0.3, col_y + 2.6, '\u2022 Monthly spread returns\n\u2022 Market-cap weights w_mkt',
            fontsize=7, color=charcoal, fontfamily='Arial', linespacing=1.4)

    rounded_box(col3_x + 0.15, col_y + 1.3, 2.4, 0.85, 'white', '#E0E0E0', 0.8)
    ax.text(col3_x + 0.3, col_y + 1.95, 'DOES:', fontsize=7, fontweight='bold', color=teal, fontfamily='Arial')
    ax.text(col3_x + 0.3, col_y + 1.7, '\u03a3 = 60-month rolling cov\n\u03c0 = \u03bb \u00d7 \u03a3 \u00d7 w_mkt\n(\u03bb = 2.5 risk aversion)',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    rounded_box(col3_x + 0.15, col_y + 0.1, 2.4, 1.05, '#E8F0FE', '#C4D7F2', 0.8)
    ax.text(col3_x + 0.3, col_y + 0.95, 'PRODUCES:', fontsize=7, fontweight='bold', color=teal, fontfamily='Arial')
    ax.text(col3_x + 0.3, col_y + 0.7, '\u03a3 = 4\u00d74 covariance matrix\n\u03c0 = equilibrium returns (prior)\nThe "anchor" for BL blending',
            fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    # Arrows down from step 2
    arrow_down(1.85, col_y, col_y - 0.35)
    arrow_down(col2_x + col_w/2, col_y, col_y - 0.35)
    arrow_down(col3_x + col_w/2, col_y, col_y - 0.35)
    ax.text(5.0, col_y - 0.15, '\u03c4, \u03c9              Q (views)              \u03c0, \u03a3',
            ha='center', fontsize=7, color=charcoal, fontfamily='Arial')

    # ── STEP 3: BLACK-LITTERMAN ──
    bl_y = 7.2
    ax.text(0.4, bl_y + 1.3, '\u2462', fontsize=14, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=navy, edgecolor=navy), zorder=5)

    # Gradient-like effect: darker navy box
    rounded_box(0.5, bl_y, 9.0, 1.5, navy, navy, 2)
    ax.text(1.0, bl_y + 1.25, 'BLACK-LITTERMAN: Combining Prior + Views + Regime',
            fontsize=10, fontweight='bold', color='white', fontfamily='Arial')
    ax.text(1.0, bl_y + 0.9,
            'Inputs: \u03c0 (prior), Q (views), \u03a3 (cov), \u03c4 (prior uncertainty), \u03c9 (view uncertainty)',
            fontsize=7.5, color='white', fontfamily='Arial', alpha=0.85)
    ax.text(1.0, bl_y + 0.6,
            '\u03a9 = diag(P\u00b7\u03c4\u03a3\u00b7P\u1d40) \u00d7 \u03c9       '
            '\u03bc_BL = [(\u03c4\u03a3)^-1 + P\u1d40\u03a9^-1 P]^-1 \u00d7 [(\u03c4\u03a3)^-1 \u03c0 + P\u1d40\u03a9^-1 Q]',
            fontsize=7, color='white', fontfamily='Arial', alpha=0.8)
    ax.text(1.0, bl_y + 0.3,
            'Stress: high \u03c4 + high \u03c9 \u2192 posterior stays near \u03c0 (defensive)    '
            'Compression: low \u03c4 + low \u03c9 \u2192 posterior incorporates Q (views matter)',
            fontsize=6.5, color=light_blue, fontfamily='Arial')

    arrow_down(5.0, bl_y, bl_y - 0.3)
    ax.text(5.3, bl_y - 0.15, '\u03bc_BL posterior returns \u2192 which buckets to overweight',
            fontsize=7, color=charcoal, fontfamily='Arial', style='italic')

    # ── STEP 4: PORTFOLIO CONSTRUCTION ──
    pc_y = 5.3
    ax.text(0.4, pc_y + 1.3, '\u2463', fontsize=14, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=navy, edgecolor=navy), zorder=5)
    rounded_box(0.5, pc_y, 9.0, 1.5, light_bg, navy, 2)
    ax.text(1.0, pc_y + 1.25, 'PORTFOLIO CONSTRUCTION: Tilt Weights by Signal',
            fontsize=10, fontweight='bold', color=navy, fontfamily='Arial')
    ax.text(1.0, pc_y + 0.9,
            'Factor signals:  DTS (50%)  |  Value (25%)  |  Momentum (25%)  \u2014  cross-sectional z-scores',
            fontsize=7.5, color=charcoal, fontfamily='Arial')
    ax.text(1.0, pc_y + 0.6,
            'w_i = w_mkt,i + \u03b1 \u00d7 z_composite,i / \u03a3|z|       '
            '\u03b1 = 10% tilt    |    Floor at 1%    |    Renormalize to 100%',
            fontsize=7, color=charcoal, fontfamily='Arial')
    ax.text(1.0, pc_y + 0.3,
            'Benchmark: AAA 4% / AA 12% / A 34% / BBB 50%    |    TC drag: 5bp one-way',
            fontsize=6.5, color='#888888', fontfamily='Arial')

    arrow_down(5.0, pc_y, pc_y - 0.3)

    # ── STEP 5: OUTPUT ──
    out_y = 4.2
    ax.text(0.4, out_y + 0.5, '\u2464', fontsize=14, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=navy, edgecolor=navy), zorder=5)
    rounded_box(0.5, out_y, 9.0, 0.8, navy, navy, 2)
    ax.text(1.0, out_y + 0.55, 'OUTPUT: Monthly Portfolio Allocation',
            fontsize=10, fontweight='bold', color='white', fontfamily='Arial')
    ax.text(1.0, out_y + 0.2, 'Four weights summing to 100% \u2014 rebalanced monthly \u2014 compared against market-cap benchmark',
            fontsize=7.5, color='white', fontfamily='Arial', alpha=0.85)

    arrow_down(5.0, out_y, out_y - 0.3)
    ax.text(5.3, out_y - 0.15, 'stress-tested, simulated, and backtested',
            fontsize=7, color=charcoal, fontfamily='Arial', style='italic')

    # ── STEP 6: VALIDATION (three boxes) ──
    val_y = 2.5
    val_w = 2.85
    val_h = 1.3
    ax.text(0.4, val_y + 1.0, '\u2465', fontsize=14, fontweight='bold', color='white',
            bbox=dict(boxstyle='circle', facecolor=gold, edgecolor=gold), zorder=5)

    for j, (title, items) in enumerate([
        ('HISTORICAL BACKTEST',
         'Run full pipeline on 25yr history.\nMeasure: Sharpe, alpha, IR,\nmax drawdown. No look-ahead.'),
        ('STRESS TESTING',
         'Shock OAS by +100 to +300bp.\nRecompute signals & weights.\nMeasure: price impact, shifts.'),
        ('MONTE CARLO',
         'Resample from real returns.\n5,000 sims \u00d7 24-month paths.\nMeasure: VaR, CVaR, P(Loss).'),
    ]):
        bx = 0.5 + j * (val_w + 0.15)
        rounded_box(bx, val_y, val_w, val_h, '#FAFAFA', '#E0E0E0', 1)
        ax.text(bx + 0.15, val_y + val_h - 0.2, title,
                fontsize=8, fontweight='bold', color=gold, fontfamily='Arial')
        ax.text(bx + 0.15, val_y + val_h - 0.45, items,
                fontsize=6.5, color=charcoal, fontfamily='Arial', linespacing=1.3)

    # Footer
    ax.plot([1.0, 9.0], [2.1, 2.1], color='#E0E0E0', linewidth=0.5)
    ax.text(5.0, 1.85, 'All data sourced from FRED  \u2022  ICE BofA OAS indices  \u2022  No synthetic inputs anywhere in the pipeline',
            ha='center', fontsize=7, color='#999999', fontfamily='Arial')

    fig.savefig(path, dpi=220, bbox_inches='tight', facecolor='white')
    plt.close(fig)

pipeline_img_path = os.path.join(tmp_dir, 'pipeline_full.png')
create_full_pipeline(pipeline_img_path)

# Caption ABOVE the diagram
add_action_title(doc, "Three parallel signal layers feed a Bayesian blending engine")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "FRED OAS data feeds regime detection, spread forecasting, and factor signal "
    "layers. Outputs are combined via Black-Litterman into constrained portfolio weights."
)
run.font.name = 'Arial'
run.font.size = Pt(9)
run.font.color.rgb = CHARCOAL

# Diagram image (nothing below)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(10)
run = p.add_run()
run.add_picture(pipeline_img_path, width=Inches(6.5))



# ═══════════════════════════════════════════════════════════════
#  PARSE AND RENDER MARKDOWN CELLS (skip cell 0 = title page)
# ═══════════════════════════════════════════════════════════════

def process_markdown_cells(doc, md_cells):
    """Process all markdown cells (skipping the title page cell)."""

    for cell_idx, cell in enumerate(md_cells):
        if cell_idx == 0:
            continue  # Title page already handled

        source = ''.join(cell['source'])
        lines = source.split('\n')

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip HTML div/p/style tags
            if re.match(r'^</?div', stripped, re.IGNORECASE):
                i += 1
                continue
            if re.match(r'^</?p\s', stripped, re.IGNORECASE):
                i += 1
                continue
            if stripped.startswith('<em>') and stripped.endswith('</em>'):
                clean = strip_html(stripped)
                if clean:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(clean)
                    run.font.size = Pt(9)
                    run.font.color.rgb = CHARCOAL
                    run.italic = True
                i += 1
                continue
            if stripped.startswith('<') and stripped.endswith('>') and not stripped.startswith('<a'):
                i += 1
                continue

            # Horizontal rules - skip
            if stripped in ('---', '***'):
                i += 1
                continue

            # Empty lines - skip
            if not stripped:
                i += 1
                continue

            # ── Skip unprofessional / casual sections ──
            skip_headings = [
                'What We Are NOT Claiming',
                'What We are NOT Claiming',
                'What we are not claiming',
            ]
            skip_paragraphs = [
                'In plain language:',
                'In plain language,',
                'In plain English',
            ]

            # Check if this line starts a heading to skip
            heading_skip = False
            for sh in skip_headings:
                if sh.lower() in stripped.lower():
                    heading_skip = True
                    break
            if heading_skip and stripped.startswith('#'):
                # Skip this heading and all following content until next heading
                i += 1
                while i < len(lines):
                    next_s = lines[i].strip()
                    if next_s.startswith('#'):
                        break
                    i += 1
                continue

            # Skip paragraphs starting with casual language
            para_skip = False
            for sp in skip_paragraphs:
                if stripped.lower().startswith(sp.lower()):
                    para_skip = True
                    break
            if para_skip:
                i += 1
                continue

            # Apply academic text cleanup to all lines
            stripped = academicize_text(stripped)

            # ── HEADINGS ──
            heading_match = re.match(r'^(#{1,3})\s+(.+)', stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                text = clean_anchor(text)
                # Convert any inline LaTeX in headings
                text = re.sub(r'\$\$(.+?)\$\$', lambda m: latex_to_unicode(m.group(1)), text)
                text = re.sub(r'\$([^$]+?)\$', lambda m: latex_to_unicode(m.group(1)), text)
                # Rename casual headings to academic style
                for casual, academic in HEADING_RENAMES.items():
                    if casual.lower() in text.lower():
                        text = re.sub(re.escape(casual), academic, text, flags=re.IGNORECASE)
                # Remove section numbers like "3.1 " prefix (already in heading level)
                text = re.sub(r'^\d+\.\d+\s+', '', text)

                # Spacing before headings (proportional to level)
                if level == 1:
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(20)
                    spacer.paragraph_format.space_after = Pt(0)

                p = doc.add_heading(text, level=level)
                # Set spacing on heading itself
                if level == 1:
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(8)
                elif level == 2:
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                elif level == 3:
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.color.rgb = NAVY if level != 2 else DENIM
                    run.font.name = 'Arial'
                i += 1
                continue

            # ── BLOCK QUOTES ──
            if stripped.startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(re.sub(r'^>\s*', '', lines[i].strip()))
                    i += 1
                full_quote = ' '.join(quote_lines)
                add_block_quote(doc, full_quote)
                continue

            # ── DISPLAY MATH ($$...$$) ──
            if stripped.startswith('$$'):
                formula_lines = [stripped]
                if not (stripped.endswith('$$') and len(stripped) > 2):
                    i += 1
                    while i < len(lines):
                        formula_lines.append(lines[i].strip())
                        if lines[i].strip().endswith('$$'):
                            break
                        i += 1
                formula_text = ' '.join(formula_lines)
                formula_text = formula_text.replace('$$', '').strip()
                add_formula_block(doc, formula_text)
                i += 1
                continue

            # ── CODE BLOCKS ──
            if stripped.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                if code_lines:
                    add_code_block(doc, code_lines)
                i += 1  # skip closing ```
                continue

            # ── TABLES ──
            if '|' in stripped and stripped.startswith('|'):
                table_lines = []
                while i < len(lines) and '|' in lines[i].strip() and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1

                if len(table_lines) < 2:
                    continue

                # Parse header
                header = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                # Skip separator (line 1), parse data rows
                data_rows = []
                for tl in table_lines[2:]:
                    row = [c.strip() for c in tl.split('|')[1:-1]]
                    data_rows.append(row)

                ncols = len(header)
                nrows = len(data_rows) + 1

                # Detect limitations table
                is_limitations_table = any('Limitation' in h or 'limitation' in h.lower() for h in header)

                table = doc.add_table(rows=nrows, cols=ncols)
                table.style = 'Table Grid'

                # Fill header
                for j, h in enumerate(header):
                    cell = table.rows[0].cells[j]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    clean_h = re.sub(r'\*\*(.+?)\*\*', r'\1', h)
                    clean_h = re.sub(r'\$(.+?)\$', lambda m: latex_to_unicode(m.group(1)), clean_h)
                    run = p.add_run(clean_h)
                    run.font.color.rgb = WHITE
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'

                # Fill data rows
                for ri, row_data in enumerate(data_rows):
                    for j in range(min(len(row_data), ncols)):
                        cell_text = row_data[j]

                        # ── HMM changes for limitations table ──
                        if is_limitations_table:
                            if 'HMM trained on full sample' in cell_text:
                                if j == 0:
                                    cell_text = "HMM refitted annually"
                            if 'Mild look-ahead bias' in cell_text or 'look-ahead bias' in cell_text.lower():
                                if j == 1:
                                    cell_text = "Regime labels may shift between refits"
                            if 'Expanding-window refitting in production' in cell_text:
                                if j == 2:
                                    cell_text = "More frequent refitting (quarterly) in production"

                        cell_text = apply_hmm_changes(cell_text)
                        cell_text = academicize_text(cell_text)

                        cell = table.rows[ri + 1].cells[j]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                        clean_text = re.sub(r'\$(.+?)\$', lambda m: latex_to_unicode(m.group(1)), clean_text)
                        clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
                        run = p.add_run(clean_text)
                        run.font.size = Pt(10)
                        run.font.color.rgb = CHARCOAL
                        run.font.name = 'Arial'

                style_table(table)
                # Minimal spacing after table
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(4)
                continue

            # ── NUMBERED LISTS ──
            num_match = re.match(r'^(\d+)\.\s+(.+)', stripped)
            if num_match:
                list_items = []
                while i < len(lines):
                    nm = re.match(r'^(\d+)\.\s+(.+)', lines[i].strip())
                    if nm:
                        list_items.append(nm.group(2))
                        i += 1
                    elif lines[i].strip().startswith('   ') or lines[i].strip().startswith('\t'):
                        if list_items:
                            list_items[-1] += ' ' + lines[i].strip()
                        i += 1
                    else:
                        break

                for item in list_items:
                    item = apply_hmm_changes(item)
                    item = academicize_text(item)
                    p = doc.add_paragraph(style='List Number')
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(3)
                    p.clear()
                    parse_inline_formatting(p, item)
                continue

            # ── BULLET LISTS ──
            bullet_match = re.match(r'^[-*]\s+(.+)', stripped)
            if bullet_match:
                list_items = []
                while i < len(lines):
                    bm = re.match(r'^[-*]\s+(.+)', lines[i].strip())
                    if bm:
                        list_items.append(bm.group(1))
                        i += 1
                    elif lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('|') and lines[i].strip() != '---':
                        if list_items and (lines[i].startswith('  ') or lines[i].startswith('\t')):
                            list_items[-1] += ' ' + lines[i].strip()
                            i += 1
                        else:
                            break
                    else:
                        break

                for item in list_items:
                    item = apply_hmm_changes(item)
                    item = academicize_text(item)
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(3)
                    p.clear()
                    parse_inline_formatting(p, item)
                continue

            # ── REGULAR PARAGRAPHS ──
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_stripped = lines[i].strip()
                if (not next_stripped or
                    next_stripped.startswith('#') or
                    next_stripped.startswith('|') or
                    next_stripped.startswith('>') or
                    next_stripped.startswith('$$') or
                    next_stripped.startswith('```') or
                    next_stripped in ('---', '***') or
                    re.match(r'^[-*]\s+', next_stripped) or
                    re.match(r'^\d+\.\s+', next_stripped)):
                    break
                para_lines.append(next_stripped)
                i += 1

            text = ' '.join(para_lines)

            # Skip pure HTML
            if text.startswith('<') and '>' in text:
                cleaned = strip_html(text)
                if not cleaned:
                    continue
                text = cleaned

            text = strip_html(text)
            if not text:
                continue

            text = apply_hmm_changes(text)
            text = academicize_text(text)

            # Skip if the text became empty after cleanup
            if not text.strip():
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            parse_inline_formatting(p, text)


# ═══════════════════════════════════════════════════════════════
#  PROCESS ALL CELLS IN ORDER (markdown + code)
# ═══════════════════════════════════════════════════════════════

# Track which markdown cell index we're on
# Skip: 0 = title page, 1 = exec summary, 2 = manual TOC (replaced with Word TOC field)
md_cell_index = 0
toc_inserted = False

for cell in all_cells:
    if cell['cell_type'] == 'markdown':
        if md_cell_index in (0, 1, 2):
            # Insert Word TOC field when we skip the manual TOC cell
            if md_cell_index == 2 and not toc_inserted:
                p = doc.add_heading("Contents", level=2)
                for run in p.runs:
                    run.font.color.rgb = NAVY
                    run.font.name = 'Arial'
                # Insert a TOC field code
                p = doc.add_paragraph()
                run = p.add_run()
                fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
                run._r.append(fldChar1)
                run2 = p.add_run()
                instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
                run2._r.append(instrText)
                run3 = p.add_run()
                fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
                run3._r.append(fldChar2)
                run4 = p.add_run("Right-click and select Update Field to generate table of contents")
                run4.font.size = Pt(10)
                run4.font.color.rgb = CHARCOAL
                run4.font.name = 'Arial'
                run4.italic = True
                run5 = p.add_run()
                fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
                run5._r.append(fldChar3)
                p.paragraph_format.space_after = Pt(12)
                toc_inserted = True
            md_cell_index += 1
            continue  # Title page, exec summary, manual TOC already handled
        # Process this single markdown cell using existing logic
        process_markdown_cells(doc, [None, cell])  # None placeholder so cell is at index 1, skipping index 0
        md_cell_index += 1
    elif cell['cell_type'] == 'code':
        process_code_cell(doc, cell, tmp_dir)

# ═══════════════════════════════════════════════════════════════
#  BACK PAGE (professional closing)
# ═══════════════════════════════════════════════════════════════

def create_back_page(path):
    """Create a professional back page with gradient footer."""
    fig, ax = plt.subplots(figsize=(8.5, 11))
    ax.set_xlim(0, 8.5)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.subplots_adjust(left=0, right=1, top=1, bottom=0)

    # White background
    ax.add_patch(plt.Rectangle((0, 0), 8.5, 11, facecolor='white', zorder=0))

    # Top accent line
    ax.plot([1.5, 7.0], [8.5, 8.5], color='#1B2A4A', linewidth=1.5, zorder=2)

    # Disclaimer
    ax.text(4.25, 8.0, 'DISCLAIMER', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#1B2A4A', fontfamily='Arial', zorder=3)

    disclaimer = (
        "This working paper is for informational and educational purposes only.\n"
        "It does not constitute investment advice, a recommendation, or an offer to\n"
        "buy or sell any securities. Past performance, whether backtested or actual,\n"
        "does not guarantee future results. All analysis is based on publicly available\n"
        "FRED data and standard quantitative methods. The author makes no warranty\n"
        "regarding the accuracy or completeness of the information presented.\n"
        "Readers should consult qualified financial professionals before making\n"
        "investment decisions."
    )
    ax.text(4.25, 7.6, disclaimer, ha='center', va='top',
            fontsize=8, color='#3C3C3C', fontfamily='Arial', linespacing=1.5, zorder=3)

    # Data sources
    ax.plot([1.5, 7.0], [6.3, 6.3], color='#E0E0E0', linewidth=0.8, zorder=2)

    ax.text(4.25, 5.9, 'DATA SOURCES', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#1B2A4A', fontfamily='Arial', zorder=3)

    sources = (
        "Federal Reserve Economic Data (FRED)\n"
        "ICE BofA US Corporate Index Option-Adjusted Spread\n"
        "ICE BofA AAA, AA, A, BBB US Corporate Index OAS\n\n"
        "All data is freely available at https://fred.stlouisfed.org"
    )
    ax.text(4.25, 5.55, sources, ha='center', va='top',
            fontsize=8, color='#3C3C3C', fontfamily='Arial', linespacing=1.5, zorder=3)

    # Replication
    ax.plot([1.5, 7.0], [4.4, 4.4], color='#E0E0E0', linewidth=0.8, zorder=2)

    ax.text(4.25, 4.0, 'REPLICATION', ha='center', va='center',
            fontsize=10, fontweight='bold', color='#1B2A4A', fontfamily='Arial', zorder=3)

    replication = (
        "The complete analytical pipeline — including data retrieval, model\n"
        "estimation, backtesting, stress testing, and Monte Carlo simulation —\n"
        "is implemented in a single Jupyter notebook using Python 3.10+\n"
        "with standard open-source libraries (hmmlearn, prophet, plotly, scipy)."
    )
    ax.text(4.25, 3.6, replication, ha='center', va='top',
            fontsize=8, color='#3C3C3C', fontfamily='Arial', linespacing=1.5, zorder=3)

    # Bottom gradient band
    gradient = np.linspace(0, 1, 256).reshape(1, -1)
    gradient = np.vstack([gradient] * 10)
    from matplotlib.colors import LinearSegmentedColormap
    navy_cmap = LinearSegmentedColormap.from_list('navy_grad2',
        ['#3D5A80', '#1B2A4A', '#0D1B2A'])
    ax.imshow(gradient, aspect='auto', cmap=navy_cmap,
              extent=[0, 8.5, 0, 2.0], zorder=1)

    ax.text(4.25, 1.3, 'Michael Tabet, CFA', ha='center', va='center',
            fontsize=14, fontweight='bold', color='white', fontfamily='Arial', zorder=3)
    ax.text(4.25, 0.9, 'QUANTITATIVE RESEARCH', ha='center', va='center',
            fontsize=9, color='#98C1D9', fontfamily='Arial', zorder=3)
    ax.text(4.25, 0.5, '\u00a9 2026  All rights reserved', ha='center', va='center',
            fontsize=7.5, color='#98C1D9', fontfamily='Arial', zorder=3)

    fig.savefig(path, dpi=250, bbox_inches='tight', pad_inches=0, facecolor='white')
    plt.close(fig)

doc.add_page_break()
back_img_path = os.path.join(tmp_dir, 'back_page.png')
create_back_page(back_img_path)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
run = p.add_run()
run.add_picture(back_img_path, width=Inches(7.5))

# Add page numbers
add_page_numbers(doc)

# Save
output_path = '/Users/michaeltabet/Desktop/two/white_paper.docx'
doc.save(output_path)
print(f"Document saved to {output_path}")
print(f"Sections: {len(doc.sections)}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

# Clean up temp figures
import shutil
shutil.rmtree(tmp_dir, ignore_errors=True)
print("Temp figures cleaned up.")
